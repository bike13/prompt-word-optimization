# -*- coding: utf-8 -*-
"""
文档处理API控制器
"""
import os
import uuid
import zipfile
import shutil
import base64
import io
import asyncio
import concurrent.futures
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Request, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
import aiofiles
import logging

# 文档处理依赖
import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import requests

# AI相关
from openai import OpenAI
from dotenv import load_dotenv

# 导入提示词配置
from prompt_doc import (
    DOCUMENT_TRANSLATION_PROMPT,
    IMAGE_TRANSLATION_IMAGE_PROMPT,
    IMAGE_TRANSLATION_TEXT_PROMPT,
    DOCUMENT_SUMMARY_PROMPT
)

from utils.logger_utils import log_api_call

load_dotenv()

logger = logging.getLogger(__name__)

router = APIRouter()


class ImageInfo:
    """图片信息类"""
    def __init__(self, image_data: bytes, position: Dict[str, Any], 
                 image_id: str = None, format: str = "png"):
        self.image_data = image_data
        self.position = position  # 包含坐标、页码等信息
        self.image_id = image_id or str(uuid.uuid4())
        self.format = format
        self.base64_data = base64.b64encode(image_data).decode('utf-8')
        self.mime_type = f"image/{format}"


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        """初始化文档处理器"""
        # 初始化OpenAI客户端
        self.openai_client = OpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            base_url=os.getenv("OPENAI_API_BASE")
        )
        self.model = os.getenv("OPENAI_API_MODEL")
        
        self.lite_text_model = os.getenv("LITE_TEXT_API_MODEL")


        self.tmag_to_tmag_model = os.getenv("TMAG_TO_TMAG_API_MODEL")

        self.tmag_to_text_model = os.getenv("VISION_API_MODEL")


        # 支持的文档格式
        self.supported_formats = {
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_docx,
            '.txt': self._read_txt,
            '.md': self._read_markdown,
            '.html': self._read_html,
            '.htm': self._read_html
        }
        
        # 并行处理配置
        self.max_workers = 10  # 并行度为10
        self.batch_size = 1    # 每次处理一张图片
    
    def get_file_extension(self, file_path: str) -> str:
        """获取文件扩展名"""
        return Path(file_path).suffix.lower()
    
    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = self.get_file_extension(file_path)
        return ext in self.supported_formats
    
    async def read_document(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """
        读取文档内容和图片
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            Tuple[文档内容, 图片信息列表]
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if not self.is_supported_format(file_path):
            raise ValueError(f"不支持的文件格式: {self.get_file_extension(file_path)}")
        
        ext = self.get_file_extension(file_path)
        read_func = self.supported_formats[ext]
                             
        try:
            content, images = await read_func(file_path)
            logger.info(f"成功读取文档: {file_path}, 内容长度: {len(content)}, 图片数量: {len(images)}")
            return content, images
        except Exception as e:
            logger.error(f"读取文档失败: {file_path}, 错误: {str(e)}")
            raise
    

    async def _read_pdf(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """读取PDF文档 - 记录图片位置信息"""
        content = ""
        images = []
        
        try:
            # 使用PyMuPDF读取PDF
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 提取文本内容
                text = page.get_text()
                
                # 提取图片并记录位置信息
                image_list = page.get_images()
                page_images = []
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # 确保不是CMYK
                            img_data = pix.tobytes("png")
                            
                            # 获取图片位置信息
                            img_rect = page.get_image_rects(xref)[0] if page.get_image_rects(xref) else None
                            position = {
                                "page": page_num + 1,
                                "index": img_index,
                                "rect": img_rect.irect if img_rect else None,
                                "xref": xref
                            }
                            
                            image_info = ImageInfo(
                                image_data=img_data,
                                position=position,
                                format="png"
                            )
                            page_images.append(image_info)
                            images.append(image_info)
                        
                        pix = None
                    except Exception as e:
                        logger.warning(f"提取PDF图片失败: {str(e)}")
                        continue
                
                # 在文本内容中插入图片位置标记
                if text.strip():
                    # 如果有图片，在文本末尾添加图片位置标记
                    if page_images:
                        for img_info in page_images:
                            # 使用图片坐标信息作为唯一标识
                            img_marker = f"image-{img_info.position['page']}-{img_info.position['index']}-{img_info.position['xref']}.png"
                            text += f"\n\n![{img_marker}]({img_marker})\n"
                    
                    content += f"--- 第 {page_num + 1} 页 ---\n\n{text.strip()}\n\n"
                else:
                    # 即使没有文本，如果有图片也要添加图片标记
                    if page_images:
                        img_markers = []
                        for img_info in page_images:
                            img_marker = f"image-{img_info.position['page']}-{img_info.position['index']}-{img_info.position['xref']}.png"
                            img_markers.append(f"![{img_marker}]({img_marker})")
                        content += f"--- 第 {page_num + 1} 页 ---\n\n" + "\n\n".join(img_markers) + "\n\n"
                    else:
                        content += f"--- 第 {page_num + 1} 页 ---\n\n[此页无文本内容]\n\n"
            
            doc.close()
            
            # 如果所有方法都失败，记录警告
            if not content.strip():
                logger.warning(f"PDF文件 {file_path} 无法提取文本内容，可能是扫描版PDF或加密PDF")
                content = f"""[警告] 无法从PDF文件中提取文本内容。"""
            
        except Exception as e:
            logger.error(f"读取PDF失败: {str(e)}")
            raise
        
        return content, images
    
    async def _read_docx(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """读取Word文档"""
        content = ""
        images = []
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
                content += "\n"
            
            # 提取图片
            # 注意：python-docx对图片提取支持有限，这里使用替代方案
            # 可以通过解压docx文件来获取图片
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                for file_info in zip_file.filelist:
                    if file_info.filename.startswith('word/media/'):
                        try:
                            img_data = zip_file.read(file_info.filename)
                            
                            # 确定图片格式
                            ext = Path(file_info.filename).suffix.lower()
                            if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                                format_name = ext[1:] if ext != '.jpeg' else 'jpg'
                                
                                position = {
                                    "filename": file_info.filename,
                                    "size": file_info.file_size
                                }
                                
                                image_info = ImageInfo(
                                    image_data=img_data,
                                    position=position,
                                    format=format_name
                                )
                                images.append(image_info)
                        except Exception as e:
                            logger.warning(f"提取Word图片失败: {str(e)}")
                            continue
            
        except Exception as e:
            logger.error(f"读取Word文档失败: {str(e)}")
            raise
        
        return content, images
    
    async def _read_txt(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """读取TXT文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, []
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
                return content, []
            except Exception as e:
                logger.error(f"读取TXT文档失败: {str(e)}")
                raise
    
    async def _read_markdown(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """读取Markdown文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, []
        except Exception as e:
            logger.error(f"读取Markdown文档失败: {str(e)}")
            raise
    
    async def _read_html(self, file_path: str) -> Tuple[str, List[ImageInfo]]:
        """读取HTML文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 提取文本内容
            content = soup.get_text()
            
            # 提取图片
            images = []
            img_tags = soup.find_all('img')
            for img_index, img_tag in enumerate(img_tags):
                src = img_tag.get('src')
                if src:
                    try:
                        # 处理相对路径和绝对路径
                        if src.startswith('http'):
                            # 网络图片
                            response = requests.get(src, timeout=10)
                            img_data = response.content
                        else:
                            # 本地图片
                            img_path = os.path.join(os.path.dirname(file_path), src)
                            if os.path.exists(img_path):
                                with open(img_path, 'rb') as f:
                                    img_data = f.read()
                            else:
                                continue
                        
                        # 确定图片格式
                        format_name = "png"  # 默认格式
                        if src.lower().endswith(('.jpg', '.jpeg')):
                            format_name = "jpg"
                        elif src.lower().endswith('.gif'):
                            format_name = "gif"
                        elif src.lower().endswith('.bmp'):
                            format_name = "bmp"
                        
                        position = {
                            "src": src,
                            "alt": img_tag.get('alt', ''),
                            "index": img_index
                        }
                        
                        image_info = ImageInfo(
                            image_data=img_data,
                            position=position,
                            format=format_name
                        )
                        images.append(image_info)
                        
                    except Exception as e:
                        logger.warning(f"提取HTML图片失败: {str(e)}")
                        continue
            
            return content, images
            
        except Exception as e:
            logger.error(f"读取HTML文档失败: {str(e)}")
            raise
    
    async def translate_document(self, content: str, target_language: str = "中文") -> str:
        """
        翻译文档内容
        
        Args:
            content: 文档内容
            target_language: 目标语言
            
        Returns:
            翻译后的内容
        """
        try:
            prompt = DOCUMENT_TRANSLATION_PROMPT.format(
                content=content,
                target_language=target_language
            )
            
            response = self.openai_client.chat.completions.create(
                model=self.lite_text_model,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            translated_content = response.choices[0].message.content
            logger.info(f"文档翻译完成，原长度: {len(content)}, 翻译后长度: {len(translated_content)}")
            return translated_content
            
        except Exception as e:
            logger.error(f"文档翻译失败: {str(e)}")
            raise
    
    async def translate_images_image_batch(self, images: List[ImageInfo], target_language: str = "中文") -> List[ImageInfo]:
        """
        并行翻译图片内容（每次处理一张图片，并行度为10）
        集成版本：所有逻辑都在此方法内，不调用其他方法
        
        Args:
            images: 图片信息列表
            target_language: 目标语言
            
        Returns:
            翻译后的图片信息列表
        """
        if not images:
            return []
        
        # 定义内部函数：翻译单张图片
        async def translate_single_image(image: ImageInfo, target_language: str) -> ImageInfo:
            """内部函数：翻译单张图片"""
            max_retries = 3
            delay = 1.0
            
            for attempt in range(max_retries):
                try:
                    # 构建消息内容 - 使用图片翻译提示词
                    messages = [
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": IMAGE_TRANSLATION_IMAGE_PROMPT.format(target_language=target_language)},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{image.mime_type};base64,{image.base64_data}"
                                    }
                                }
                            ]
                        }
                    ]
                    
                    # 使用线程池执行API调用
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                        future = executor.submit(
                            self.openai_client.chat.completions.create,
                            model=self.tmag_to_tmag_model,
                            messages=messages,
                            temperature=0.3
                        )
                        response = future.result()
                        translated_text = response.choices[0].message.content
                        
                        # 创建翻译后的图片信息（这里简化处理，使用原图数据）
                        # 在实际应用中，这里应该根据translated_text生成新的图片
                        translated_image = ImageInfo(
                            image_data=image.image_data,
                            position=image.position,
                            image_id=f"{image.image_id}_translated",
                            format=image.format
                        )
                        return translated_image
                        
                except Exception as e:
                    error_msg = str(e)
                    logger.warning(f"翻译图片失败 (尝试 {attempt + 1}/{max_retries}): {error_msg}")
                    
                    # 检查是否是可重试的错误
                    retryable_errors = ['500', 'InternalServiceError', 'rate_limit', 'timeout', 'InternalServerError']
                    if any(error in error_msg for error in retryable_errors) and attempt < max_retries - 1:
                        wait_time = delay * (2 ** attempt)  # 指数退避
                        logger.info(f"等待 {wait_time:.1f} 秒后重试...")
                        await asyncio.sleep(wait_time)
                        continue
                    
                    # 不可重试的错误或达到最大重试次数
                    logger.error(f"翻译图片最终失败: {error_msg}")
                    raise e
            
            # 如果所有重试都失败了，返回原图
            logger.warning(f"图片翻译失败，返回原图")
            return image

        
        
        # 创建任务列表，每张图片一个任务
        tasks = []
        for image in images:
            task = translate_single_image(image, target_language)
            tasks.append(task)
        
        # 并行执行所有任务，最大并行度为10
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 处理异常结果
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"翻译第{i+1}张图片失败: {str(result)}")
                processed_results.append(f"翻译失败: {str(result)}")
            else:
                processed_results.append(result)
        
        return processed_results
    
    async def translate_images_text_batch(self, images: List[ImageInfo], target_language: str = "中文") -> List[str]:
        """
        并行翻译图片内容（每次处理一张图片，并行度为10）
        集成版本：所有逻辑都在此方法内，不调用其他方法
        
        Args:
            images: 图片信息列表
            target_language: 目标语言
            
        Returns:
            翻译结果列表
        """
        if not images:
            return []
        
        # 定义内部函数：翻译单张图片
        async def translate_single_image(image: ImageInfo, target_language: str) -> str:
            """内部函数：翻译单张图片"""
            try:
                # 构建消息内容 - 使用图片翻译提示词
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": IMAGE_TRANSLATION_TEXT_PROMPT.format(target_language=target_language)},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{image.mime_type};base64,{image.base64_data}"
                                }
                            }
                        ]
                    }
                ]
                
                # 使用线程池执行同步API调用
                loop = asyncio.get_event_loop()
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(
                        self.openai_client.chat.completions.create,
                        model=self.tmag_to_text_model,
                        messages=messages,
                        temperature=0.3
                    )
                    response = await loop.run_in_executor(None, future.result)
                    return response.choices[0].message.content
                    
            except Exception as e:
                logger.error(f"翻译图片失败: {str(e)}")
                raise e

        # 使用信号量控制并发度
        semaphore = asyncio.Semaphore(10)
        
        async def translate_with_semaphore(image: ImageInfo, target_language: str) -> str:
            async with semaphore:
                return await translate_single_image(image, target_language)
        
        # 创建任务列表，每张图片一个任务
        tasks = []
        for image in images:
            task = translate_with_semaphore(image, target_language)
            tasks.append(task)
        
        # 并行执行所有任务，最大并行度为10
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 处理异常结果
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"翻译第{i+1}张图片失败: {str(result)}")
                processed_results.append(f"翻译失败: {str(result)}")
            else:
                processed_results.append(result)
        
        return processed_results

    
    async def summarize_document(self, content: str, images: List[ImageInfo]) -> str:
        """
        总结文档内容
        
        Args:
            content: 文档内容
            images: 图片信息列表
            
        Returns:
            总结内容
        """
        try:
            # 构建包含图片的提示词
            image_context = ""
            if images:
                image_context = f"\n\n文档包含 {len(images)} 张图片，请结合图片内容进行总结。"
            
            prompt = DOCUMENT_SUMMARY_PROMPT.format(content=content) + image_context
            
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                reasoning_effort="medium"
            )
            
            summary = response.choices[0].message.content
            logger.info(f"文档总结完成，原长度: {len(content)}, 总结长度: {len(summary)}")
            return summary
            
        except Exception as e:
            logger.error(f"文档总结失败: {str(e)}")
            raise
    
    def combine_document_with_images(self, original_content: str, translated_content: str, 
                                   original_images: List[ImageInfo], 
                                   image_translations: List[str],
                                   output_dir: str) -> str:
        """
        组合拼接文档和图片，将图片翻译文字放在原图下方
        
        Args:
            original_content: 原始文档内容
            translated_content: 翻译后的文档内容（可能包含图片标记）
            original_images: 原始图片列表
            image_translations: 图片翻译文字列表
            output_dir: 输出目录
            
        Returns:
            拼接后的Markdown内容
            
        Note:
            - 原始内容部分：显示原图 + 翻译文字
            - 翻译内容部分：用原图 + 翻译文字替换原有的图片位置标记
        """
        try:
            # 创建assets目录
            assets_dir = os.path.join(output_dir, "assets")
            os.makedirs(assets_dir, exist_ok=True)
            
            # 创建图片映射字典，用于快速查找图片
            image_map = {}
            for i, (orig_img, trans_text) in enumerate(zip(original_images, image_translations)):
                # 使用位置信息创建唯一标识
                img_key = f"image-{orig_img.position['page']}-{orig_img.position['index']}-{orig_img.position['xref']}.png"
                
                # 保存原始图片
                orig_filename = f"{img_key.replace('.png', '')}-original.{orig_img.format}"
                orig_path = os.path.join(assets_dir, orig_filename)
                with open(orig_path, 'wb') as f:
                    f.write(orig_img.image_data)
                
                # 存储图片映射信息（译文部分使用原图+翻译文字）
                image_map[img_key] = {
                    'original': orig_filename,
                    'translation': trans_text,
                    'position': orig_img.position
                }
            
            # 替换原文中的图片引用，将图片和翻译文字组合显示
            processed_original_content = original_content
            for img_key, img_info in image_map.items():
                # 替换原文中的图片引用为图片+翻译文字的组合
                original_pattern = f"![{img_key}]({img_key})"
                original_replacement = f"""![{img_key.replace('.png', '-original')}](assets/{img_info['original']})"""
                processed_original_content = processed_original_content.replace(original_pattern, original_replacement)
            
            # 替换译文中的图片引用：用原图+翻译文字替换原有的图片位置标记
            processed_translated_content = translated_content
            for img_key, img_info in image_map.items():
                # 在翻译内容中，如果存在图片标记，用原图+翻译文字的组合替换它
                translated_pattern = f"![{img_key}]({img_key})"
                translated_replacement = f"""![{img_key.replace('.png', '-original')}](assets/{img_info['original']})

**图片翻译：** 
{img_info['translation']}"""
                processed_translated_content = processed_translated_content.replace(translated_pattern, translated_replacement)

            # 构建Markdown内容
            markdown_content = f"""
## 原始内容
{processed_original_content}

## 翻译内容
{processed_translated_content}

"""
            
            return markdown_content
            
        except Exception as e:
            logger.error(f"组合文档失败: {str(e)}")
            raise


# 初始化文档处理器
doc_processor = DocumentProcessor()

# 数据目录
DATA_DIR = "data"
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")


def ensure_directories():
    """确保必要的目录存在"""
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(UPLOAD_DIR, exist_ok=True)


def get_device_dir(device_id: str) -> str:
    """获取设备目录"""
    device_dir = os.path.join(UPLOAD_DIR, device_id)
    os.makedirs(device_dir, exist_ok=True)
    return device_dir


@router.post("/upload")
@log_api_call('/api/doc/upload', 'POST')
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    device_id: str = Form(default="default")
):
    """
    上传文件接口
    
    Args:
        file: 上传的文件
        device_id: 设备ID，用于区分不同设备的文件
        
    Returns:
        上传结果信息
    """
    try:
        ensure_directories()
        
        # 检查文件格式
        if not doc_processor.is_supported_format(file.filename):
            raise HTTPException(
                status_code=400, 
                detail=f"不支持的文件格式: {Path(file.filename).suffix}"
            )
        
        # 获取设备目录
        device_dir = get_device_dir(device_id)
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_ext = Path(file.filename).suffix
        new_filename = f"{file_id}{file_ext}"
        file_path = os.path.join(device_dir, new_filename)
        
        # 保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # 读取文档内容用于预览
        try:
            content, images = await doc_processor.read_document(file_path)
            
            return JSONResponse(content={
                "success": True,
                "message": "文件上传成功",
                "data": {
                    "file_id": file_id,
                    "filename": file.filename,
                    "file_path": file_path,
                    "device_id": device_id,
                    "content_preview": content[:1000] + "..." if len(content) > 1000 else content,
                    "image_count": len(images),
                    "file_size": len(content)
                }
            })
            
        except Exception as e:
            logger.error(f"读取上传文件失败: {str(e)}")
            # 即使读取失败，文件也已上传成功
            return JSONResponse(content={
                "success": True,
                "message": "文件上传成功，但读取内容失败",
                "data": {
                    "file_id": file_id,
                    "filename": file.filename,
                    "file_path": file_path,
                    "device_id": device_id,
                    "error": str(e)
                }
            })
            
    except Exception as e:
        logger.error(f"文件上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")


@router.post("/translate")
@log_api_call('/api/doc/translate', 'POST')
async def translate_document(
    request: Request,
    file_id: str = Form(...),
    device_id: str = Form(default="default"),
    target_language: str = Form(default="中文")
):
    """
    AI翻译接口
    
    Args:
        file_id: 文件ID
        device_id: 设备ID
        target_language: 目标语言
        
    Returns:
        翻译结果
    """
    try:
        device_dir = get_device_dir(device_id)
        
        # 查找文件
        file_path = None
        for file in os.listdir(device_dir):
            if file.startswith(file_id):
                file_path = os.path.join(device_dir, file)
                break
        
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取文档内容
        content, images = await doc_processor.read_document(file_path)
        
        # 翻译文档内容
        translated_content = await doc_processor.translate_document(content, target_language)
        
        # 翻译图片内容为文字和图片
        image_translations = []
        translated_images = []
        if images:
            # 获取图片翻译文字
            image_translations = await doc_processor.translate_images_text_batch(images, target_language)
            # 获取翻译后的图片
            # translated_images = await doc_processor.translate_images_image_batch(images, target_language)
        
        # 组合拼接文档
        original_filename = Path(file_path).stem
        translated_filename = f"{original_filename}(译文).md"
        translated_file_path = os.path.join(device_dir, translated_filename)
        
        combined_content = doc_processor.combine_document_with_images(
            content, translated_content, images, image_translations,  device_dir
        )
        
        # 保存翻译结果
        async with aiofiles.open(translated_file_path, 'w', encoding='utf-8') as f:
            await f.write(combined_content)
        
        return JSONResponse(content={
            "success": True,
            "message": "翻译完成",
            "data": {
                "file_id": file_id,
                "translated_content": combined_content,
                "translated_file_path": translated_file_path,
                "original_length": len(content),
                "translated_length": len(translated_content),
                "image_count": len(images)
            }
        })
        
    except Exception as e:
        logger.error(f"文档翻译失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"翻译失败: {str(e)}")


@router.post("/summarize")
@log_api_call('/api/doc/summarize', 'POST')
async def summarize_document(
    request: Request,
    file_id: str = Form(...),
    device_id: str = Form(default="default")
):
    """
    AI总结接口
    
    Args:
        file_id: 文件ID
        device_id: 设备ID
        
    Returns:
        总结结果
    """
    try:
        device_dir = get_device_dir(device_id)
        
        # 查找文件
        file_path = None
        for file in os.listdir(device_dir):
            if file.startswith(file_id):
                file_path = os.path.join(device_dir, file)
                break
        
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取文档内容
        content, images = await doc_processor.read_document(file_path)
        
        # 总结文档内容
        summary = await doc_processor.summarize_document(content, images)
        
        # 保存总结结果
        original_filename = Path(file_path).stem
        summary_filename = f"{original_filename}(总结).md"
        summary_file_path = os.path.join(device_dir, summary_filename)
        
        async with aiofiles.open(summary_file_path, 'w', encoding='utf-8') as f:
            await f.write(summary)
        
        return JSONResponse(content={
            "success": True,
            "message": "总结完成",
            "data": {
                "file_id": file_id,
                "summary": summary,
                "summary_file_path": summary_file_path,
                "original_length": len(content),
                "summary_length": len(summary),
                "image_count": len(images)
            }
        })
        
    except Exception as e:
        logger.error(f"文档总结失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"总结失败: {str(e)}")


@router.get("/download/{device_id}")
@log_api_call('/api/doc/download', 'GET')
async def download_files(device_id: str, background_tasks: BackgroundTasks):
    """
    下载文件接口
    
    Args:
        device_id: 设备ID
        
    Returns:
        ZIP文件下载
    """
    try:
        device_dir = get_device_dir(device_id)
        
        if not os.path.exists(device_dir):
            raise HTTPException(status_code=404, detail="设备目录不存在")
        
        # 检查目录是否为空
        files = os.listdir(device_dir)
        if not files:
            raise HTTPException(status_code=404, detail="没有文件可下载")
        
        # 创建ZIP文件
        zip_filename = f"{device_id}_files.zip"
        zip_path = os.path.join(DATA_DIR, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(device_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, device_dir)
                    zipf.write(file_path, arcname)
        
        # 定义清理函数
        def cleanup_files():
            """删除临时文件"""
            try:
                # 延迟删除，确保文件下载完成
                import time
                time.sleep(2)
                
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                    logger.info(f"已删除临时ZIP文件: {zip_path}")
                
                # 删除设备目录
                if os.path.exists(device_dir):
                    shutil.rmtree(device_dir)
                    logger.info(f"已删除设备目录: {device_dir}")
                    
            except Exception as e:
                logger.error(f"清理临时文件失败: {str(e)}")
        
        # 添加后台任务，在响应发送后清理文件
        background_tasks.add_task(cleanup_files)
        
        # 使用FileResponse返回文件
        response = FileResponse(
            path=zip_path,
            filename=zip_filename,
            media_type='application/zip'
        )
        
        return response
        
    except Exception as e:
        logger.error(f"文件下载失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载失败: {str(e)}")


@router.get("/list/{device_id}")
@log_api_call('/api/doc/list', 'GET')
async def list_files(device_id: str):
    """
    列出设备文件
    
    Args:
        device_id: 设备ID
        
    Returns:
        文件列表
    """
    try:
        device_dir = get_device_dir(device_id)
        
        if not os.path.exists(device_dir):
            return JSONResponse(content={
                "success": True,
                "data": {
                    "files": [],
                    "device_id": device_id
                }
            })
        
        files = []
        for file in os.listdir(device_dir):
            file_path = os.path.join(device_dir, file)
            if os.path.isfile(file_path):
                stat = os.stat(file_path)
                files.append({
                    "filename": file,
                    "size": stat.st_size,
                    "modified": stat.st_mtime,
                    "file_id": file.split('.')[0] if '.' in file else file
                })
        
        return JSONResponse(content={
            "success": True,
            "data": {
                "files": files,
                "device_id": device_id
            }
        })
        
    except Exception as e:
        logger.error(f"列出文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"列出文件失败: {str(e)}")


@router.delete("/delete/{device_id}/{file_id}")
@log_api_call('/api/doc/delete', 'DELETE')
async def delete_file(device_id: str, file_id: str):
    """
    删除文件
    
    Args:
        device_id: 设备ID
        file_id: 文件ID
        
    Returns:
        删除结果
    """
    try:
        device_dir = get_device_dir(device_id)
        
        # 查找并删除文件
        deleted_files = []
        for file in os.listdir(device_dir):
            if file.startswith(file_id):
                file_path = os.path.join(device_dir, file)
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    deleted_files.append(file)
        
        if not deleted_files:
            raise HTTPException(status_code=404, detail="文件不存在")
        
        return JSONResponse(content={
            "success": True,
            "message": "文件删除成功",
            "data": {
                "deleted_files": deleted_files,
                "device_id": device_id,
                "file_id": file_id
            }
        })
        
    except Exception as e:
        logger.error(f"删除文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")
